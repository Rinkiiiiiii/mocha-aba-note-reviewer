"""
File -> plain text extraction for PDF and Word (.docx) source documents.

Design note: position/layout on the page is intentionally discarded here.
Downstream code detection (see `core.detection`) works by searching the
resulting plain text for known CPT code tokens, which — as confirmed against
real sample notes from multiple different EHR/practice-management vendors —
is robust across wildly different field layouts, whereas trying to parse
*where* a code sits on the page is not.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Union

import docx  # python-docx
from pypdf import PdfReader


class UnsupportedFileTypeError(ValueError):
    """Raised for file types this extractor doesn't (yet) support."""


@dataclass
class ExtractedDocument:
    full_text: str
    page_texts: list[str] = field(default_factory=list)
    docusign_stripped: bool = False
    warnings: list[str] = field(default_factory=list)


# --- DocuSign "Certificate of Completion" stripping --------------------------
#
# Every DocuSign-signed sample we reviewed appends 1-2 pages of audit-trail
# boilerplate (envelope IDs, IP addresses, signature images, consent
# disclosures) after the actual clinical note. It carries zero clinical
# content and would waste prompt space / distract the AI reviewer, so we cut
# it before it ever reaches detection or the AI client.
#
# We require the primary marker PLUS at least one confirming signal so we
# never accidentally truncate a real note that happens to mention
# "certificate of completion" in passing (unlikely, but cheap to guard).

_DOCUSIGN_MARKER = re.compile(r"certificate of completion", re.IGNORECASE)
_DOCUSIGN_CONFIRMATION_SIGNALS = (
    "envelope id",
    "signer events",
    "security level:",
    "consent of electronic record",
    "esigned",
)


def _strip_docusign_boilerplate(text: str) -> tuple[str, bool]:
    match = _DOCUSIGN_MARKER.search(text)
    if not match:
        return text, False

    tail_lower = text[match.start():].lower()
    hits = sum(1 for signal in _DOCUSIGN_CONFIRMATION_SIGNALS if signal in tail_lower)
    if hits >= 2:
        return text[: match.start()].rstrip(), True
    return text, False


# --- PDF ----------------------------------------------------------------

def _extract_pdf(path: Path) -> tuple[list[str], list[str]]:
    warnings: list[str] = []
    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # pypdf can raise a variety of parsing errors
            warnings.append(f"Failed to extract text from page {i + 1}: {exc}")
            pages.append("")

    if not any(p.strip() for p in pages):
        warnings.append(
            "No extractable text found in this PDF — it may be a scanned/"
            "image-only document. OCR is not currently supported."
        )
    return pages, warnings


# --- DOCX -----------------------------------------------------------------

def _extract_docx(path: Path) -> tuple[list[str], list[str]]:
    warnings: list[str] = []
    document = docx.Document(str(path))

    blocks: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            blocks.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    full = "\n".join(blocks)
    if not full.strip():
        warnings.append("No extractable text found in this document.")

    # python-docx has no reliable concept of page boundaries, so the whole
    # document is treated as a single "page" for downstream purposes.
    return [full], warnings


# --- Public entry point ----------------------------------------------------

def extract_document_text(path: Union[str, Path]) -> ExtractedDocument:
    """Extract plain text from a PDF or Word (.docx) file, with DocuSign
    audit-trail boilerplate stripped when detected.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        page_texts, warnings = _extract_pdf(path)
    elif suffix == ".docx":
        page_texts, warnings = _extract_docx(path)
    elif suffix == ".doc":
        raise UnsupportedFileTypeError(
            "Legacy .doc files aren't supported — please save/export as .docx or PDF."
        )
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {suffix or '(none)'}")

    full_text = "\n".join(page_texts)
    stripped_text, was_stripped = _strip_docusign_boilerplate(full_text)

    return ExtractedDocument(
        full_text=stripped_text,
        page_texts=page_texts,
        docusign_stripped=was_stripped,
        warnings=warnings,
    )
