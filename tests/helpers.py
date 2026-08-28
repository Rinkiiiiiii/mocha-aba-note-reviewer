"""
Test-only helpers for generating synthetic (non-PHI) PDF/DOCX fixtures.

These are built programmatically rather than committed as binary files so
the repo never carries any client data — everything here uses invented
names/dates that mimic the *structural* patterns observed in real samples
(see architecture-decisions.md) without containing anything real.
"""

from __future__ import annotations

from pathlib import Path

import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def make_pdf(path: Path, pages: list[list[str]]) -> None:
    """Create a simple text PDF. `pages` is a list of pages, each a list of
    lines to draw top-to-bottom on that page.
    """
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    for page_lines in pages:
        y = height - 72
        for line in page_lines:
            c.drawString(72, y, line)
            y -= 14
            if y < 72:
                c.showPage()
                y = height - 72
        c.showPage()
    c.save()


def make_docx(path: Path, paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> None:
    document = docx.Document()
    for para in paragraphs:
        document.add_paragraph(para)
    for table_data in tables or []:
        rows = len(table_data)
        cols = len(table_data[0]) if rows else 0
        table = document.add_table(rows=rows, cols=cols)
        for r, row_data in enumerate(table_data):
            for c_idx, cell_text in enumerate(row_data):
                table.cell(r, c_idx).text = cell_text
    document.save(str(path))


DOCUSIGN_BOILERPLATE_LINES = [
    "Certificate of Completion",
    "Envelope Id: 00000000-0000-0000-0000-000000000000",
    "Status: Completed",
    "Subject: Document",
    "Signer Events                    Signature",
    "Test Signer                      ESIGNED",
    "testsigner@example.com",
    "Security Level: Session, Account Authentication",
    "Consent of Electronic Record and Signature Disclosure Used:",
    "Using IP Address: 10.0.0.1",
    "Signed on: 01/01/2026 at 9:00am EDT",
]
